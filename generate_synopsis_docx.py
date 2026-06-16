from __future__ import annotations

import os
import re
from pathlib import Path

BASE_DIR = Path(__file__).resolve().parent
MPL_CONFIG_DIR = BASE_DIR / ".mplconfig"
MPL_CONFIG_DIR.mkdir(parents=True, exist_ok=True)
os.environ.setdefault("MPLCONFIGDIR", str(MPL_CONFIG_DIR))

import matplotlib.pyplot as plt
from matplotlib.patches import Ellipse, FancyBboxPatch, Polygon
from docx import Document
from docx.shared import Inches


INPUT_MD = BASE_DIR / "SYNOPSIS_2.md"
OUTPUT_DOCX = BASE_DIR / "SYNOPSIS_2_WELL_EXPLAINED.docx"
DIAGRAM_DIR = BASE_DIR / "synopsis_diagrams"


def _setup_ax(title: str, width: float = 14, height: float = 8):
    fig, ax = plt.subplots(figsize=(width, height))
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 100)
    ax.axis("off")
    ax.set_title(title, fontsize=16, weight="bold", pad=12)
    return fig, ax


def _box(ax, x, y, w, h, text, fc="#f7fbff", ec="#1f4e79", fontsize=9):
    patch = FancyBboxPatch(
        (x, y), w, h, boxstyle="round,pad=0.3", linewidth=1.5, edgecolor=ec, facecolor=fc
    )
    ax.add_patch(patch)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fontsize, wrap=True)


def _ellipse(ax, x, y, w, h, text, fc="#fff7e6", ec="#7f6000", fontsize=9):
    e = Ellipse((x, y), w, h, linewidth=1.5, edgecolor=ec, facecolor=fc)
    ax.add_patch(e)
    ax.text(x, y, text, ha="center", va="center", fontsize=fontsize, wrap=True)


def _diamond(ax, x, y, w, h, text, fc="#fde9d9", ec="#833c0c", fontsize=9):
    points = [(x, y + h / 2), (x + w / 2, y + h), (x + w, y + h / 2), (x + w / 2, y)]
    d = Polygon(points, closed=True, linewidth=1.5, edgecolor=ec, facecolor=fc)
    ax.add_patch(d)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fontsize, wrap=True)


def _arrow(ax, x1, y1, x2, y2, label=None, fs=8):
    ax.annotate(
        "",
        xy=(x2, y2),
        xytext=(x1, y1),
        arrowprops=dict(arrowstyle="->", lw=1.4, color="#333333"),
    )
    if label:
        mx = (x1 + x2) / 2
        my = (y1 + y2) / 2
        ax.text(mx, my + 1.8, label, fontsize=fs, ha="center", va="center")


def _poly_arrow(ax, points, label=None, fs=8):
    if len(points) < 2:
        return
    for i in range(len(points) - 2):
        x1, y1 = points[i]
        x2, y2 = points[i + 1]
        ax.plot([x1, x2], [y1, y2], color="#333333", lw=1.2)
    x1, y1 = points[-2]
    x2, y2 = points[-1]
    _arrow(ax, x1, y1, x2, y2)
    if label:
        lx = sum(p[0] for p in points) / len(points)
        ly = sum(p[1] for p in points) / len(points)
        ax.text(lx, ly + 1.6, label, fontsize=fs, ha="center", va="center")


def _save(fig, path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def draw_architecture(path: Path):
    fig, ax = _setup_ax("System Architecture")
    _box(ax, 8, 62, 20, 14, "User Browser")
    _box(ax, 8, 28, 20, 14, "Admin Browser")
    _box(ax, 38, 42, 24, 18, "Flask Web\nApplication", fc="#eaf3ff")
    _box(ax, 72, 65, 20, 14, "Database\n(MySQL/SQLite)", fc="#eef7ea", ec="#2e7d32")
    _box(ax, 72, 42, 20, 14, "Payment Layer\n(Simulated)", fc="#fff3e8", ec="#a15c00")
    _box(ax, 72, 19, 20, 14, "Chat Assistant\n(DB + AI + Fallback)", fc="#f3e8ff", ec="#5a2d82")
    _arrow(ax, 28, 69, 38, 51)
    _arrow(ax, 28, 35, 38, 49)
    _arrow(ax, 62, 54, 72, 72)
    _arrow(ax, 62, 51, 72, 49)
    _arrow(ax, 62, 45, 72, 26)
    _arrow(ax, 72, 68, 62, 56)
    _arrow(ax, 72, 49, 62, 49)
    _arrow(ax, 72, 26, 62, 46)
    _save(fig, path)


def draw_dfd_level0(path: Path):
    fig, ax = _setup_ax("DFD Level 0 - Context Diagram")
    _box(ax, 8, 58, 20, 14, "User")
    _box(ax, 8, 28, 20, 14, "Admin")
    _box(ax, 38, 41, 26, 18, "Holiday Booking\nSystem", fc="#eaf3ff")
    _box(ax, 74, 58, 18, 14, "Payment\nGateway")
    _box(ax, 74, 28, 18, 14, "Database", fc="#eef7ea", ec="#2e7d32")
    _arrow(ax, 28, 65, 38, 54, "Search/Book/Pay")
    _arrow(ax, 38, 47, 28, 61, "Results/Status")
    _arrow(ax, 28, 35, 38, 47, "Manage Data")
    _arrow(ax, 38, 44, 28, 32, "Admin Views")
    _arrow(ax, 64, 56, 74, 65, "Payment Req")
    _arrow(ax, 74, 61, 64, 53, "Response")
    _arrow(ax, 64, 46, 74, 35, "CRUD")
    _arrow(ax, 74, 31, 64, 44, "Records")
    _save(fig, path)


def draw_dfd_level1(path: Path):
    fig, ax = _setup_ax("DFD Level 1 - Major Processes", width=16, height=9)
    _box(ax, 4, 70, 14, 12, "User")
    _box(ax, 4, 20, 14, 12, "Admin")
    _box(ax, 24, 76, 18, 10, "P1 Auth")
    _box(ax, 24, 61, 18, 10, "P2 Package\nSearch")
    _box(ax, 24, 46, 18, 10, "P3 Package\nBooking")
    _box(ax, 24, 31, 18, 10, "P4 Transport\nBooking")
    _box(ax, 24, 16, 18, 10, "P5 Wishlist,\nReview, Contact")
    _box(ax, 24, 1, 18, 10, "P6 Admin\nManagement")

    _box(ax, 56, 78, 14, 10, "D1 Users", fc="#eef7ea", ec="#2e7d32")
    _box(ax, 56, 64, 14, 10, "D2 Packages", fc="#eef7ea", ec="#2e7d32")
    _box(ax, 56, 50, 14, 10, "D3 Package\nBookings", fc="#eef7ea", ec="#2e7d32")
    _box(ax, 56, 36, 14, 10, "D4 Transport\nMasters", fc="#eef7ea", ec="#2e7d32")
    _box(ax, 56, 22, 14, 10, "D5 Transport\nBookings", fc="#eef7ea", ec="#2e7d32")
    _box(ax, 56, 8, 14, 10, "D6 Reviews,\nWishlist,\nContacts", fc="#eef7ea", ec="#2e7d32")
    _box(ax, 78, 45, 16, 12, "Payment\nLayer")

    for y in (80, 65, 50, 35, 20):
        _arrow(ax, 18, 76, 24, y)
    _arrow(ax, 18, 26, 24, 6)

    _arrow(ax, 42, 81, 56, 83)
    _arrow(ax, 42, 66, 56, 69)
    _arrow(ax, 42, 51, 56, 55)
    _arrow(ax, 42, 36, 56, 41)
    _arrow(ax, 42, 21, 56, 27)
    _arrow(ax, 42, 6, 56, 13)
    _arrow(ax, 70, 55, 78, 51)
    _arrow(ax, 78, 49, 70, 53)
    _arrow(ax, 42, 51, 78, 49, "payment")
    _arrow(ax, 42, 36, 78, 49, "payment")
    _save(fig, path)


def draw_dfd_level2(path: Path):
    fig, ax = _setup_ax("DFD Level 2 - Package Booking Process", width=10, height=12)
    _box(ax, 30, 90, 40, 7, "Start: Select Package")
    _diamond(ax, 30, 80, 40, 8, "Logged in?")
    _box(ax, 30, 69, 40, 7, "Redirect to Login")
    _box(ax, 30, 58, 40, 7, "Enter Travellers")
    _box(ax, 30, 47, 40, 7, "Calculate Total Amount")
    _box(ax, 30, 36, 40, 7, "Enter Card Details")
    _diamond(ax, 30, 25, 40, 8, "Card fields valid?")
    _box(ax, 30, 14, 40, 7, "Create Booking Record")
    _box(ax, 30, 4, 40, 7, "Show Success in Dashboard")
    _box(ax, 74, 25, 20, 8, "Show Payment\nError")

    _arrow(ax, 50, 90, 50, 88)
    _arrow(ax, 50, 80, 50, 76, "No")
    _arrow(ax, 50, 69, 50, 65)
    _arrow(ax, 50, 80, 50, 65, "Yes")
    _arrow(ax, 50, 58, 50, 54)
    _arrow(ax, 50, 47, 50, 43)
    _arrow(ax, 50, 36, 50, 33)
    _arrow(ax, 50, 25, 50, 21, "Yes")
    _arrow(ax, 50, 14, 50, 11)
    _arrow(ax, 70, 29, 74, 29, "No")
    _save(fig, path)


def draw_er(path: Path):
    fig, ax = _setup_ax("ER Diagram - Holiday Booking System", width=18, height=11)

    # Left column: user/profile interactions
    _box(ax, 4, 82, 18, 10, "USER\n(id, username,\nemail, is_admin)")
    _box(ax, 4, 64, 18, 10, "PACKAGE\n(id, name, price,\nduration)")
    _box(ax, 4, 46, 18, 10, "REVIEW\n(package_id, user_id,\nrating)")
    _box(ax, 4, 28, 18, 10, "WISHLIST\n(package_id, user_id)")
    _box(ax, 4, 10, 18, 10, "CONTACT_US\n(name, email,\nmessage, status)")

    # Middle column: booking + transport masters
    _box(ax, 36, 82, 18, 10, "BOOKING\n(user_id, package_id,\nstatus, total_price)")
    _box(ax, 36, 64, 18, 10, "FLIGHT\n(airline, route,\nbase_price)")
    _box(ax, 36, 46, 18, 10, "TRAIN\n(operator, route,\nbase_price)")
    _box(ax, 36, 28, 18, 10, "BUS\n(operator, route,\nbase_price)")
    _box(ax, 36, 10, 18, 10, "CAB\n(provider, route,\nbase_price)")

    # Right column: transport booking tables
    _box(ax, 68, 64, 18, 10, "FLIGHT_BOOKING\n(flight_id, user_id,\nclass, final_price)")
    _box(ax, 68, 46, 18, 10, "TRAIN_BOOKING\n(train_id, user_id,\nclass, final_price)")
    _box(ax, 68, 28, 18, 10, "BUS_BOOKING\n(bus_id, user_id,\nclass, final_price)")
    _box(ax, 68, 10, 18, 10, "CAB_BOOKING\n(cab_id, user_id,\nclass, final_price)")

    # Package/user core relationships
    _arrow(ax, 22, 87, 36, 87, "1:N")  # USER -> BOOKING
    _arrow(ax, 22, 69, 36, 85, "1:N")  # PACKAGE -> BOOKING
    _poly_arrow(ax, [(22, 87), (26, 87), (26, 51), (22, 51)], "1:N")  # USER -> REVIEW
    _poly_arrow(ax, [(22, 69), (30, 69), (30, 51), (22, 51)], "1:N")  # PACKAGE -> REVIEW
    _poly_arrow(ax, [(22, 87), (26, 87), (26, 33), (22, 33)], "1:N")  # USER -> WISHLIST
    _poly_arrow(ax, [(22, 69), (32, 69), (32, 33), (22, 33)], "1:N")  # PACKAGE -> WISHLIST
    _poly_arrow(ax, [(22, 87), (28, 87), (28, 15), (22, 15)], "1:N")  # USER -> CONTACT_US

    # Transport master -> transport booking
    _arrow(ax, 54, 69, 68, 69, "1:N")  # FLIGHT -> FLIGHT_BOOKING
    _arrow(ax, 54, 51, 68, 51, "1:N")  # TRAIN -> TRAIN_BOOKING
    _arrow(ax, 54, 33, 68, 33, "1:N")  # BUS -> BUS_BOOKING
    _arrow(ax, 54, 15, 68, 15, "1:N")  # CAB -> CAB_BOOKING

    # USER -> transport booking links (routed through a clear corridor)
    _poly_arrow(ax, [(22, 87), (60, 87), (60, 69), (68, 69)], "1:N")
    _poly_arrow(ax, [(22, 87), (60, 87), (60, 51), (68, 51)], "1:N")
    _poly_arrow(ax, [(22, 87), (60, 87), (60, 33), (68, 33)], "1:N")
    _poly_arrow(ax, [(22, 87), (60, 87), (60, 15), (68, 15)], "1:N")

    # Optional booking link to flight booking
    _poly_arrow(ax, [(54, 87), (64, 87), (64, 69), (68, 69)], "0..1:1")

    _save(fig, path)


def draw_use_case(path: Path):
    fig, ax = _setup_ax("Use Case Diagram", width=16, height=9)
    _box(ax, 6, 64, 12, 16, "User")
    _box(ax, 82, 64, 12, 16, "Admin")
    _box(ax, 82, 20, 12, 16, "Payment\nLayer")

    _ellipse(ax, 32, 77, 24, 10, "Register / Login")
    _ellipse(ax, 32, 63, 24, 10, "Browse and Search")
    _ellipse(ax, 32, 49, 24, 10, "Book Package")
    _ellipse(ax, 32, 35, 24, 10, "Book Transport")
    _ellipse(ax, 32, 21, 24, 10, "Wishlist / Review /\nContact")

    _ellipse(ax, 62, 77, 24, 10, "Admin Login")
    _ellipse(ax, 62, 63, 24, 10, "Manage Packages")
    _ellipse(ax, 62, 49, 24, 10, "Manage Transport")
    _ellipse(ax, 62, 35, 24, 10, "Edit Bookings")
    _ellipse(ax, 62, 21, 24, 10, "Handle Inquiries")
    _ellipse(ax, 47, 8, 24, 10, "Make Payment")

    for y in (77, 63, 49, 35, 21):
        _arrow(ax, 18, 72, 20, y)
    for y in (77, 63, 49, 35, 21):
        _arrow(ax, 82, 72, 74, y)
    _arrow(ax, 47, 13, 82, 28)
    _arrow(ax, 32, 49, 47, 13)
    _arrow(ax, 32, 35, 47, 13)
    _save(fig, path)


def draw_flow_auth(path: Path):
    fig, ax = _setup_ax("Flowchart - User Registration and Login", width=10, height=12)
    _box(ax, 28, 90, 44, 7, "Open Register/Login")
    _diamond(ax, 32, 80, 36, 8, "New User?")
    _box(ax, 5, 69, 40, 7, "Fill Registration Form")
    _diamond(ax, 5, 58, 40, 8, "Email exists?")
    _box(ax, 5, 47, 40, 7, "Create Account")
    _box(ax, 55, 69, 40, 7, "Open Login Form")
    _diamond(ax, 55, 58, 40, 8, "Credentials Valid?")
    _box(ax, 55, 47, 40, 7, "Login Success")
    _box(ax, 55, 36, 40, 7, "Go to Dashboard")
    _box(ax, 5, 36, 40, 7, "Show Error")

    _arrow(ax, 50, 90, 50, 88)
    _arrow(ax, 43, 84, 25, 73, "Yes")
    _arrow(ax, 57, 84, 75, 73, "No")
    _arrow(ax, 25, 69, 25, 66)
    _arrow(ax, 25, 58, 25, 54, "No")
    _arrow(ax, 25, 58, 25, 39, "Yes")
    _arrow(ax, 25, 47, 75, 73)
    _arrow(ax, 75, 69, 75, 66)
    _arrow(ax, 75, 58, 75, 54, "Yes")
    _arrow(ax, 75, 58, 25, 39, "No")
    _arrow(ax, 75, 47, 75, 43)
    _save(fig, path)


def draw_flow_package(path: Path):
    fig, ax = _setup_ax("Flowchart - Package Booking and Payment", width=10, height=12)
    _box(ax, 28, 90, 44, 7, "Select Package")
    _box(ax, 28, 79, 44, 7, "Open Checkout")
    _box(ax, 28, 68, 44, 7, "Enter Members")
    _box(ax, 28, 57, 44, 7, "Calculate Total Price")
    _box(ax, 28, 46, 44, 7, "Enter Card Details")
    _diamond(ax, 28, 35, 44, 8, "Card data valid?")
    _box(ax, 28, 24, 44, 7, "Create Booking")
    _box(ax, 28, 13, 44, 7, "Save in Database")
    _box(ax, 28, 2, 44, 7, "Show Confirmation")
    _box(ax, 76, 35, 20, 8, "Show Error")

    for top, bottom in [(90, 86), (79, 75), (68, 64), (57, 53), (46, 43), (24, 20), (13, 9)]:
        _arrow(ax, 50, top, 50, bottom)
    _arrow(ax, 50, 35, 50, 31, "Yes")
    _arrow(ax, 72, 39, 76, 39, "No")
    _save(fig, path)


def draw_flow_transport(path: Path):
    fig, ax = _setup_ax("Flowchart - Transport Booking", width=10, height=12)
    _box(ax, 24, 90, 52, 7, "Open Flight/Train/Bus/Cab Page")
    _box(ax, 24, 79, 52, 7, "Enter Route, Date, Class, Travellers")
    _box(ax, 24, 68, 52, 7, "Fetch Matching Options")
    _box(ax, 24, 57, 52, 7, "Compute Per-Person and Total Fare")
    _box(ax, 24, 46, 52, 7, "Open Checkout and Confirm")
    _diamond(ax, 24, 35, 52, 8, "Validation Success?")
    _box(ax, 24, 24, 52, 7, "Create Transport Booking")
    _box(ax, 24, 13, 52, 7, "Save and Redirect Dashboard")
    _box(ax, 80, 35, 18, 8, "Return\nCheckout")

    for top, bottom in [(90, 86), (79, 75), (68, 64), (57, 53), (46, 43), (24, 20)]:
        _arrow(ax, 50, top, 50, bottom)
    _arrow(ax, 50, 35, 50, 31, "Yes")
    _arrow(ax, 76, 39, 80, 39, "No")
    _arrow(ax, 50, 13, 50, 9)
    _save(fig, path)


def draw_flow_admin(path: Path):
    fig, ax = _setup_ax("Flowchart - Admin Operations", width=10, height=12)
    _box(ax, 28, 90, 44, 7, "Admin Login")
    _box(ax, 28, 79, 44, 7, "Open Admin Dashboard")
    _diamond(ax, 28, 68, 44, 8, "Select Operation")
    _box(ax, 6, 57, 38, 7, "Manage Packages")
    _box(ax, 56, 57, 38, 7, "Manage Transport")
    _box(ax, 6, 46, 38, 7, "Edit Bookings")
    _box(ax, 56, 46, 38, 7, "Handle Inquiries")
    _box(ax, 28, 33, 44, 7, "Commit DB Updates")
    _box(ax, 28, 22, 44, 7, "Show Success/Error")

    _arrow(ax, 50, 90, 50, 86)
    _arrow(ax, 50, 79, 50, 76)
    _arrow(ax, 50, 68, 25, 60)
    _arrow(ax, 50, 68, 75, 60)
    _arrow(ax, 50, 68, 25, 49)
    _arrow(ax, 50, 68, 75, 49)
    _arrow(ax, 25, 57, 50, 36)
    _arrow(ax, 75, 57, 50, 36)
    _arrow(ax, 25, 46, 50, 36)
    _arrow(ax, 75, 46, 50, 36)
    _arrow(ax, 50, 33, 50, 29)
    _save(fig, path)


def draw_flow_chatbot(path: Path):
    fig, ax = _setup_ax("Flowchart - Chatbot Response", width=10, height=12)
    _box(ax, 28, 90, 44, 7, "User Sends Message")
    _diamond(ax, 28, 79, 44, 8, "Message Empty?")
    _box(ax, 4, 68, 40, 7, "Ask User to Enter Query")
    _box(ax, 54, 68, 42, 7, "Try DB-Aware Response")
    _diamond(ax, 54, 57, 42, 8, "DB Response Found?")
    _box(ax, 54, 46, 42, 7, "Return DB Reply")
    _box(ax, 4, 46, 40, 7, "Try AI Response")
    _diamond(ax, 4, 35, 40, 8, "AI Available?")
    _box(ax, 4, 24, 40, 7, "Return AI Reply")
    _box(ax, 54, 24, 42, 7, "Return Fallback Reply")

    _arrow(ax, 50, 90, 50, 87)
    _arrow(ax, 28, 83, 24, 71, "Yes")
    _arrow(ax, 50, 79, 75, 71, "No")
    _arrow(ax, 75, 68, 75, 65)
    _arrow(ax, 75, 57, 75, 53, "Yes")
    _arrow(ax, 54, 61, 24, 49, "No")
    _arrow(ax, 24, 46, 24, 43)
    _arrow(ax, 24, 35, 24, 31, "Yes")
    _arrow(ax, 44, 39, 54, 27, "No")
    _save(fig, path)


def _draw_all_diagrams() -> dict[str, Path]:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    files = {
        "architecture": DIAGRAM_DIR / "architecture.png",
        "dfd_0": DIAGRAM_DIR / "dfd_level_0.png",
        "dfd_1": DIAGRAM_DIR / "dfd_level_1.png",
        "dfd_2": DIAGRAM_DIR / "dfd_level_2.png",
        "er": DIAGRAM_DIR / "er_diagram.png",
        "use_case": DIAGRAM_DIR / "use_case_diagram.png",
        "flow_auth": DIAGRAM_DIR / "flow_auth.png",
        "flow_package": DIAGRAM_DIR / "flow_package.png",
        "flow_transport": DIAGRAM_DIR / "flow_transport.png",
        "flow_admin": DIAGRAM_DIR / "flow_admin.png",
        "flow_chatbot": DIAGRAM_DIR / "flow_chatbot.png",
    }
    draw_architecture(files["architecture"])
    draw_dfd_level0(files["dfd_0"])
    draw_dfd_level1(files["dfd_1"])
    draw_dfd_level2(files["dfd_2"])
    draw_er(files["er"])
    draw_use_case(files["use_case"])
    draw_flow_auth(files["flow_auth"])
    draw_flow_package(files["flow_package"])
    draw_flow_transport(files["flow_transport"])
    draw_flow_admin(files["flow_admin"])
    draw_flow_chatbot(files["flow_chatbot"])
    return files


def _add_markdown_table(doc: Document, table_lines: list[str]):
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if cells:
            rows.append(cells)
    if len(rows) < 2:
        for line in table_lines:
            doc.add_paragraph(line)
        return
    header = rows[0]
    body = rows[2:] if len(rows) > 2 else []
    table = doc.add_table(rows=1 + len(body), cols=len(header))
    table.style = "Table Grid"
    for i, c in enumerate(header):
        table.cell(0, i).text = c
    for r_idx, row in enumerate(body, start=1):
        padded = row + [""] * (len(header) - len(row))
        for c_idx, c in enumerate(padded[: len(header)]):
            table.cell(r_idx, c_idx).text = c


def build_docx(diagrams: dict[str, Path]):
    if not INPUT_MD.exists():
        raise FileNotFoundError(f"Missing source file: {INPUT_MD}")

    image_map = {
        "## 8. System Architecture": diagrams["architecture"],
        "### 12.1 DFD Level 0 (Context Diagram)": diagrams["dfd_0"],
        "### 12.2 DFD Level 1 (Major Processes)": diagrams["dfd_1"],
        "### 12.3 DFD Level 2 (Package Booking Process)": diagrams["dfd_2"],
        "## 13. Entity Relationship (ER) Diagram - Detailed": diagrams["er"],
        "## 14. Use Case Diagram": diagrams["use_case"],
        "### 15.1 User Registration and Login Flow": diagrams["flow_auth"],
        "### 15.2 Package Booking and Payment Flow": diagrams["flow_package"],
        "### 15.3 Transport Booking Flow (Flight/Train/Bus/Cab)": diagrams["flow_transport"],
        "### 15.4 Admin Data Management Flow": diagrams["flow_admin"],
        "### 15.5 Chatbot Response Flow": diagrams["flow_chatbot"],
    }

    doc = Document()
    md_lines = INPUT_MD.read_text(encoding="utf-8").splitlines()

    in_code = False
    table_buffer: list[str] = []

    for line in md_lines:
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            continue

        if in_code:
            continue

        if stripped.startswith("|"):
            table_buffer.append(line)
            continue
        if table_buffer:
            _add_markdown_table(doc, table_buffer)
            table_buffer = []

        if not stripped:
            doc.add_paragraph("")
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = min(len(m.group(1)), 4)
            title = m.group(2).strip()
            key = f"{m.group(1)} {title}"
            doc.add_heading(title, level=level)
            if key in image_map:
                doc.add_picture(str(image_map[key]), width=Inches(6.7))
            continue

        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue

        if re.match(r"^\d+\.\s+", stripped):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
            continue

        doc.add_paragraph(line)

    if table_buffer:
        _add_markdown_table(doc, table_buffer)

    doc.save(OUTPUT_DOCX)


def main():
    diagrams = _draw_all_diagrams()
    build_docx(diagrams)
    print(f"Created: {OUTPUT_DOCX}")
    print(f"Diagrams folder: {DIAGRAM_DIR}")


if __name__ == "__main__":
    main()
